"""Tailorer: generates a per-job tailored .docx, stores in your schema.

Writes:
- `resumes` row (if base doesn't exist yet) — represents the candidate's base resume
- `resume_versions` row — the per-job tailored version, with file_url pointing
  to either Supabase Storage (if upload succeeded) or the local path

Output:
- ~/Downloads/{Company}_{Role}_{date}.docx
- + Optionally uploaded to Supabase Storage bucket
"""
from __future__ import annotations

import base64
import io
import json
import logging
import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from daemon import config, llm
from daemon.supabase_client import (
    create_tailored_version,
    log_event,
    storage,
    upsert_base_resume,
    upsert_job_status,
)


log = logging.getLogger(__name__)


SYSTEM_PROMPT = """You are an expert resume writer. You tailor a candidate's existing resume for a specific job.

CRITICAL RULES:
1. NEVER invent facts. Only re-order, rephrase, and re-emphasize real experience.
2. Keep the resume ONE PAGE. ATS-friendly. No tables/graphics/columns.
3. Lead the summary with 1-2 keywords from the JD that the candidate has real experience with.
4. Skills section must list the candidate's actual skills, but reorder so the JD's required skills come FIRST.
5. Work history: keep all real bullets. Re-order WITHIN each role so JD-matching bullets come first.
6. If the candidate is missing something the JD asks for, do NOT pretend they have it.

Return STRICT JSON only:
{
  "summary": "3-line professional summary rewritten for this role",
  "skills_top": ["top 10 actual skills, reordered"],
  "work_history": [
    {"company": "...", "title": "...", "dates": "...", "bullets": ["..."]}
  ],
  "education": [{"school": "...", "degree": "...", "year": "..."}],
  "tailoring_notes": "1-3 line human-readable summary of what you changed and why",
  "keywords_surfaced": ["keyword1", "keyword2", ...]
}
"""


def _safe_json_parse(raw: str) -> dict[str, Any]:
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", (raw or "").strip(), flags=re.MULTILINE)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass
        log.error("tailor JSON parse failed; raw=%s", (raw or "")[:500])
        return {}


def _safe_filename(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9_-]+", "_", s).strip("_")
    return s[:80] or "resume"


def render_docx(payload: dict[str, Any], *, output_path: Path, profile: dict[str, Any]) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(36)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    full_name = profile.get("full_name", "")
    contact_bits = [
        profile.get("email"),
        profile.get("phone"),
        profile.get("location"),
        profile.get("linkedin_url"),
    ]
    contact = " | ".join(b for b in contact_bits if b)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(full_name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    if contact:
        c = doc.add_paragraph(contact)
        for r in c.runs:
            r.font.size = Pt(9)

    summary = payload.get("summary") or profile.get("summary", "")
    if summary:
        h = doc.add_paragraph()
        hr = h.add_run("Summary")
        hr.bold = True
        hr.font.size = Pt(11)
        doc.add_paragraph(summary)

    skills_top = payload.get("skills_top") or profile.get("skills", [])
    if skills_top:
        h = doc.add_paragraph()
        hr = h.add_run("Skills")
        hr.bold = True
        hr.font.size = Pt(11)
        doc.add_paragraph(" · ".join(skills_top))

    work_history = payload.get("work_history") or profile.get("work_history", [])
    if work_history:
        h = doc.add_paragraph()
        hr = h.add_run("Experience")
        hr.bold = True
        hr.font.size = Pt(11)
        for role in work_history:
            company = role.get("company", "")
            title = role.get("title", "")
            dates = role.get("dates", "")
            line = f"{title} — {company}"
            if dates:
                line += f" ({dates})"
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(10.5)
            for bullet in role.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(bullet).font.size = Pt(10.5)

    education = payload.get("education") or profile.get("education", [])
    if education:
        h = doc.add_paragraph()
        hr = h.add_run("Education")
        hr.bold = True
        hr.font.size = Pt(11)
        for ed in education:
            doc.add_paragraph(f"{ed.get('degree','')} — {ed.get('school','')} ({ed.get('year','')})")

    doc.save(str(output_path))


def _upload_to_storage(local_path: Path, job_id: str) -> str | None:
    """Try to upload the .docx to Supabase Storage. Returns the PUBLIC URL or None.

    Storage layout: {bucket}/{job_id}/{filename}
    Public URL form: {SUPABASE_URL}/storage/v1/object/public/{bucket}/{job_id}/{filename}
    """
    if not config.RESUMES_BUCKET:
        return None
    try:
        bucket = storage().from_(config.RESUMES_BUCKET)
        storage_path = f"{job_id}/{local_path.name}"
        with open(local_path, "rb") as f:
            data = f.read()
        # Remove any existing file at this path so upload succeeds (idempotent overwrite)
        try:
            bucket.remove([storage_path])
        except Exception:
            pass
        bucket.upload(
            storage_path,
            data,
            {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
        )
        public_url = f"{config.SUPABASE_URL}/storage/v1/object/public/{config.RESUMES_BUCKET}/{storage_path}"
        log.info("Uploaded to storage: %s", public_url)
        return public_url
    except Exception as e:
        log.warning("Storage upload failed (will fall back to local file_url): %s", e)
        return None


def tailor_resume(job: dict, profile: dict, keywords: list[str] | None = None) -> dict[str, Any]:
    """Tailor resume for this job. Writes to resumes + resume_versions. Returns summary."""
    job_id = job["id"]
    desc_text = job.get("description") or ""
    if not desc_text:
        return {"ok": False, "error": "no description"}

    upsert_job_status(job_id, "tailoring")
    log_event(job_id, "tailoring_started")

    user_prompt = (
        f"Candidate profile (JSON):\n{json.dumps(profile, indent=2, default=str)}\n\n"
        f"Job description:\n{desc_text[:6000]}\n\n"
        f"Top keywords to surface:\n{json.dumps(keywords or [])}\n\n"
        f"Return strict JSON only."
    )

    try:
        resp = llm.call(
            model=config.CLAUDE_MODEL_TAILOR,
            system=SYSTEM_PROMPT,
            user=user_prompt,
            max_tokens=4096,
        )
    except llm.BudgetExceeded as e:
        return {"ok": False, "error": str(e)}
    except Exception as e:
        log.exception("tailor call failed: %s", e)
        return {"ok": False, "error": str(e)}

    payload = _safe_json_parse(resp["text"])
    if not payload:
        return {"ok": False, "error": "empty parser output"}

    # Filename: {Company}_{Role}_{date}.docx
    company = (job.get("company_name") or job.get("_job_company") or "Resume").strip()
    role = (job.get("job_title") or job.get("_job_role") or "Tailored").strip()
    safe_company = _safe_filename(company)
    safe_role = _safe_filename(role)
    filename = f"{safe_company}_{safe_role}_{date.today().isoformat()}.docx"
    output_path = config.DOWNLOADS_DIR / filename

    try:
        render_docx(payload, output_path=output_path, profile=profile)
    except Exception as e:
        log.exception("docx render failed: %s", e)
        return {"ok": False, "error": f"render failed: {e}"}

    # Upload to Supabase Storage (best-effort)
    storage_path = _upload_to_storage(output_path, job_id)
    file_url = storage_path or str(output_path)

    # Ensure base resume row exists (idempotent)
    base_resume_id = upsert_base_resume(
        profile_id=str(profile.get("id") or ""),
        file_url=str(output_path),  # local path of original base resume
        version_label="base",
    )

    # Insert resume_versions
    version_id = create_tailored_version(
        job_id=job_id,
        base_resume_id=base_resume_id,
        tailored_resume_path=str(output_path),
        tailored_resume_url=storage_path or str(output_path),
        tailoring_notes=payload.get("tailoring_notes", ""),
    )

    upsert_job_status(job_id, "tailored", message=payload.get("tailoring_notes", "")[:300])
    log_event(job_id, "tailored", {
        "file_path": str(output_path),
        "storage_path": storage_path,
        "keywords_surfaced": payload.get("keywords_surfaced", []),
        "llm_cost_usd": resp["cost_usd"],
    })

    return {
        "ok": True,
        "file_path": str(output_path),
        "file_name": filename,
        "storage_path": storage_path,
        "storage_uploaded": storage_path is not None,
        "version_id": version_id,
        "tailoring_notes": payload.get("tailoring_notes", ""),
        "keywords_surfaced": payload.get("keywords_surfaced", []),
        "llm_cost_usd": resp["cost_usd"],
    }